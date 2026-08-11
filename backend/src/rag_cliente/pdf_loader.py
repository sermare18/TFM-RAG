"""Carga y extracción de documentos para el índice RAG.

Este módulo convierte documentos locales en una lista homogénea de
`PageDocument`, que es la estructura que consume el chunker del proyecto.

Formatos soportados:
- Con Marker 2 full: PDF, DOCX, PPTX, XLSX, EPUB, HTML e imágenes.
- Texto plano (`.txt`): lectura UTF-8.
- Si desactivo Marker, conservo fallbacks nativos para PDF digital, DOCX y TXT.

Decisión principal:
Uso exclusivamente el converter y los builders oficiales de Marker 2. La
salida primaria es JSON estructurado; el renderer Markdown anterior queda
disponible solo mediante un flag temporal de compatibilidad.
"""
from __future__ import annotations

import html as html_lib
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

import fitz
from docx import Document

from rag_cliente.config import ResolvedMarkerProfile, Settings, resolve_marker_profile
from rag_cliente.marker_capabilities import (
    marker_capabilities,
    require_marker_installed_and_warn_if_unvalidated,
)

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
MARKER_DOCUMENT_SUFFIXES = {
    ".pdf",
    ".docx",
    ".pptx",
    ".xlsx",
    ".epub",
    ".html",
    *IMAGE_SUFFIXES,
}
SUPPORTED_DOCUMENT_SUFFIXES = {".txt", *MARKER_DOCUMENT_SUFFIXES}
ProgressCallback = Callable[[str], None]
_MARKER_PAGE_SEPARATOR_RE = re.compile(
    r"(?:^|\n+)(?:\{)?(\d+)(?:\})?-{20,}[ \t]*(?:\n+|$)"
)

# Pipeline estándar de Marker 2.0.0 sin procesadores LLM ajenos a tablas. Los
# perfiles quality insertan los dos procesadores oficiales inmediatamente
# después de TableProcessor. cpu-digital conserva el mismo pipeline no LLM.
MARKER_STANDARD_NON_LLM_PROCESSORS: tuple[str, ...] = (
    "marker.processors.block_relabel.BlockRelabelProcessor",
    "marker.processors.line_merge.LineMergeProcessor",
    "marker.processors.blockquote.BlockquoteProcessor",
    "marker.processors.code.CodeProcessor",
    "marker.processors.document_toc.DocumentTOCProcessor",
    "marker.processors.equation.EquationProcessor",
    "marker.processors.footnote.FootnoteProcessor",
    "marker.processors.ignoretext.IgnoreTextProcessor",
    "marker.processors.line_numbers.LineNumbersProcessor",
    "marker.processors.list.ListProcessor",
    "marker.processors.page_header.PageHeaderProcessor",
    "marker.processors.marginalia.MarginaliaProcessor",
    "marker.processors.sectionheader.SectionHeaderProcessor",
    "marker.processors.table.TableProcessor",
    "marker.processors.text.TextProcessor",
    "marker.processors.reference.ReferenceProcessor",
    "marker.processors.blank_page.BlankPageProcessor",
    "marker.processors.debug.DebugProcessor",
)
MARKER_OFFICIAL_TABLE_LLM_PROCESSORS: tuple[str, ...] = (
    "marker.processors.llm.llm_table.LLMTableProcessor",
    "marker.processors.llm.llm_table_merge.LLMTableMergeProcessor",
)


@dataclass(slots=True)
class DocumentElement:
    """Elemento estructurado normalizado desde el JSON oficial de Marker."""

    id: str
    kind: str
    html: str
    text: str
    page_start: int
    page_end: int
    source_pages: list[int] = field(default_factory=list)
    source_block_ids: list[str] = field(default_factory=list)
    source_spans: list[dict[str, Any]] = field(default_factory=list)
    section_hierarchy: dict[str, Any] = field(default_factory=dict)
    polygon: list[list[float]] = field(default_factory=list)
    confidence: float | None = None
    children: list["DocumentElement"] = field(default_factory=list)
    document_id: str = ""
    parent_id: str | None = None
    table_id: str | None = None
    section_path: list[str] = field(default_factory=list)
    provenance: str = "marker"
    metadata: dict[str, Any] = field(default_factory=dict)

    def as_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "kind": self.kind,
            "html": self.html,
            "text": self.text,
            "page_start": self.page_start,
            "page_end": self.page_end,
            "source_pages": list(self.source_pages),
            "source_block_ids": list(self.source_block_ids),
            "source_spans": [dict(span) for span in self.source_spans],
            "section_hierarchy": dict(self.section_hierarchy),
            "polygon": [list(point) for point in self.polygon],
            "confidence": self.confidence,
            "children": [child.as_dict() for child in self.children],
            "document_id": self.document_id,
            "parent_id": self.parent_id,
            "table_id": self.table_id,
            "section_path": list(self.section_path),
            "provenance": self.provenance,
            "metadata": dict(self.metadata),
        }


@dataclass(slots=True)
class ParsedDocument:
    """Resultado primario del parser, independiente del chunking posterior."""

    id: str
    source: str
    source_path: str
    source_type: str
    elements: list[DocumentElement]
    metadata: dict[str, Any] = field(default_factory=dict)

    def as_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "source": self.source,
            "source_path": self.source_path,
            "source_type": self.source_type,
            "elements": [element.as_dict() for element in self.elements],
            "metadata": dict(self.metadata),
        }


@dataclass(slots=True)
class PageDocument:
    """Unidad de contenido ya extraída y lista para chunking.

    Atributos:
        document_id: Identificador lógico del documento.
        source: Nombre visible del archivo original.
        source_path: Ruta absoluta al archivo original.
        source_type: Tipo de documento: pdf, docx, txt o extensión de imagen.
        page_number: Página/unidad lógica asociada al texto.
        text: Contenido textual normalizado.
        ocr_used: Indica si el texto de esta página/unidad procede de OCR.
        tag: Etiqueta derivada de la primera carpeta relativa bajo el directorio indexado.
        block_type/id/html/page/polygon/children/section_hierarchy: Estructura
            oficial conservada desde la salida JSON de Marker.
        extraction_metadata: Metadatos de extracción del documento y la página.
    """

    document_id: str
    source: str
    source_path: str
    source_type: str
    page_number: int
    text: str
    ocr_used: bool = False
    tag: str = ""
    block_type: str = ""
    id: str = ""
    html: str = ""
    page: int | None = None
    polygon: list[list[float]] = field(default_factory=list)
    children: list[dict[str, Any]] = field(default_factory=list)
    section_hierarchy: dict[str, Any] = field(default_factory=dict)
    page_start: int | None = None
    page_end: int | None = None
    source_pages: list[int] = field(default_factory=list)
    source_block_ids: list[str] = field(default_factory=list)
    source_spans: list[dict[str, Any]] = field(default_factory=list)
    confidence: float | None = None
    extraction_metadata: dict[str, Any] = field(default_factory=dict)
    provenance: str = "marker"
    parser_profile: str = ""

    def __post_init__(self) -> None:
        if self.page_start is None:
            self.page_start = self.page_number
        if self.page_end is None:
            self.page_end = self.page_number
        if not self.source_pages:
            self.source_pages = list(range(self.page_start, self.page_end + 1))


def _emit(progress_callback: ProgressCallback | None, message: str) -> None:
    if progress_callback is not None:
        progress_callback(message)


def _normalize_text(text: str) -> str:
    """Normaliza bloques de texto preservando separación por párrafos."""
    return "\n\n".join(block.strip() for block in text.split("\n\n") if block.strip()).strip()


def _section_path_from_block(
    block: dict[str, Any],
    inherited: list[str] | None = None,
) -> list[str]:
    """Conserva la ruta de sección publicada por Marker sin inferirla del texto."""
    raw_path = block.get("section_path")
    if isinstance(raw_path, list):
        path = [str(item).strip() for item in raw_path if str(item).strip()]
        if path:
            return path

    hierarchy = _to_plain_data(block.get("section_hierarchy") or {})
    if isinstance(hierarchy, dict) and hierarchy:
        def sort_key(item: tuple[Any, Any]) -> tuple[int, str]:
            key = str(item[0])
            try:
                return int(key), key
            except ValueError:
                return 10_000, key

        path = [
            str(value).strip()
            for _, value in sorted(hierarchy.items(), key=sort_key)
            if str(value).strip()
        ]
        if path:
            return path

    return list(inherited or [])


def _extract_docx_text(document: Document) -> str:
    """Extrae párrafos y tablas de un documento Word (.docx)."""
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return _normalize_text("\n\n".join(blocks))


def _as_1_based_page_number(raw_page_id: Any) -> int | None:
    """Convierte IDs internos de Marker a numeración humana 1-based."""
    try:
        page_id = int(raw_page_id)
    except (TypeError, ValueError):
        return None
    return page_id + 1 if page_id >= 0 else None


def _collect_marker_page_extraction_details(page: Any, document: Any) -> dict[str, Any]:
    """Recojo señales internas de Marker 2 sobre OCR/visión por página.

    No me basta con `page.text_extraction_method`: en modo fast Marker puede
    reparar un bloque aislado, y en ambos modos una tabla digital de baja
    confianza puede usar el fallback visual sin convertir toda la página.
    """
    methods: set[str] = set()

    def add_method(value: Any) -> None:
        method = str(value or "").strip().lower()
        if method:
            methods.add(method)

    add_method(getattr(page, "text_extraction_method", None))

    try:
        blocks = list(page.contained_blocks(document))
    except Exception:
        blocks = []

    for block in blocks:
        add_method(getattr(block, "text_extraction_method", None))

    ocr_errors_detected = bool(getattr(page, "ocr_errors_detected", False))
    visual_methods = sorted(method for method in methods if method != "pdftext")
    ocr_used = ocr_errors_detected or bool(visual_methods)

    reasons: list[str] = []
    if ocr_errors_detected:
        reasons.append("embedded_text_rejected")
    if visual_methods:
        reasons.append("visual_block_or_table_fallback")

    return {
        "text_extraction_methods": sorted(methods),
        "ocr_errors_detected": ocr_errors_detected,
        "ocr_used": ocr_used,
        "ocr_reasons": reasons,
    }


def _metadata_with_marker_capabilities(metadata: Any) -> dict[str, Any]:
    """Añade el contrato de capacidad a los metadatos del documento."""
    enriched = dict(metadata or {}) if isinstance(metadata, dict) else {}
    enriched["capabilities"] = marker_capabilities()
    return enriched


class MarkerMarkdownRendererWithOcrMetadata:
    """Enriquezco el renderer Markdown de Marker con metadatos OCR por página.

    Delego el Markdown al renderer oficial para no duplicar su lógica y, antes
    de devolverlo, registro también reparaciones visuales parciales y tablas que
    hayan caído al OCR adaptativo de Marker 2.
    """

    def __init__(self, config: dict[str, Any] | None = None) -> None:
        from marker.renderers.markdown import MarkdownRenderer

        self._renderer = MarkdownRenderer(config=config)

    def __call__(self, document: Any) -> Any:
        rendered = self._renderer(document)
        metadata = _metadata_with_marker_capabilities(
            getattr(rendered, "metadata", {})
        )
        page_stats_by_id: dict[int, dict[str, Any]] = {}

        for page_stat in metadata.get("page_stats", []) or []:
            if not isinstance(page_stat, dict):
                continue
            try:
                page_stats_by_id[int(page_stat.get("page_id"))] = dict(page_stat)
            except (TypeError, ValueError):
                continue

        enriched_page_stats: list[dict[str, Any]] = []
        for page in getattr(document, "pages", []) or []:
            raw_page_id = getattr(page, "page_id", None)
            try:
                page_id = int(raw_page_id)
            except (TypeError, ValueError):
                continue

            page_stat = page_stats_by_id.get(page_id, {"page_id": page_id})
            extraction_details = _collect_marker_page_extraction_details(page, document)
            page_stat.update(extraction_details)
            enriched_page_stats.append(page_stat)

        if enriched_page_stats:
            metadata["page_stats"] = enriched_page_stats

        if hasattr(rendered, "model_copy"):
            return rendered.model_copy(update={"metadata": metadata})

        rendered.metadata = metadata
        return rendered


class MarkerJSONRendererWithOcrMetadata:
    """Delega la estructura al renderer JSON público y añade metadatos OCR."""

    def __init__(self, config: dict[str, Any] | None = None) -> None:
        from marker.renderers.json import JSONRenderer

        self._renderer = JSONRenderer(config=config)

    def __call__(self, document: Any) -> Any:
        rendered = self._renderer(document)
        metadata = _metadata_with_marker_capabilities(
            getattr(rendered, "metadata", {})
        )
        page_stats_by_id: dict[int, dict[str, Any]] = {}

        for page_stat in metadata.get("page_stats", []) or []:
            if not isinstance(page_stat, dict):
                continue
            try:
                page_stats_by_id[int(page_stat.get("page_id"))] = dict(page_stat)
            except (TypeError, ValueError):
                continue

        enriched_page_stats: list[dict[str, Any]] = []
        for page in getattr(document, "pages", []) or []:
            try:
                page_id = int(getattr(page, "page_id", None))
            except (TypeError, ValueError):
                continue
            page_stat = page_stats_by_id.get(page_id, {"page_id": page_id})
            page_stat.update(_collect_marker_page_extraction_details(page, document))
            enriched_page_stats.append(page_stat)

        if enriched_page_stats:
            metadata["page_stats"] = enriched_page_stats

        if hasattr(rendered, "model_copy"):
            return rendered.model_copy(update={"metadata": metadata})
        rendered.metadata = metadata
        return rendered


def _load_native_pdf_pages(pdf_path: Path) -> list[PageDocument]:
    """Fallback rápido para PDFs digitales cuando Marker está deshabilitado."""
    pages: list[PageDocument] = []
    with fitz.open(pdf_path) as document:
        for page_index, page in enumerate(document):
            text = page.get_text("text").strip()
            if not text:
                continue
            pages.append(
                PageDocument(
                    document_id=pdf_path.stem,
                    source=pdf_path.name,
                    source_path=str(pdf_path.resolve()),
                    source_type="pdf",
                    page_number=page_index + 1,
                    text=_normalize_text(text),
                    ocr_used=False,
                    provenance="pymupdf",
                    parser_profile="native-pdf",
                )
            )
    return pages


def _official_marker_processor_paths(
    profile: ResolvedMarkerProfile,
) -> tuple[str, ...]:
    """Conserva el pipeline estándar y limita LLM a las tablas oficiales."""
    processors = list(MARKER_STANDARD_NON_LLM_PROCESSORS)
    if profile.use_llm:
        table_index = processors.index("marker.processors.table.TableProcessor") + 1
        processors[table_index:table_index] = MARKER_OFFICIAL_TABLE_LLM_PROCESSORS
    return tuple(processors)


def _build_marker_config(
    settings: Settings,
    profile: ResolvedMarkerProfile | None = None,
) -> dict[str, Any]:
    """Construye configuración exacta para uno de los perfiles soportados."""
    resolved_profile = profile or resolve_marker_profile(settings)
    config: dict[str, Any] = {
        "output_format": "markdown" if settings.marker_markdown_compatibility else "json",
        "paginate_output": settings.marker_markdown_compatibility,
        "disable_image_extraction": settings.marker_disable_image_extraction,
        "mode": resolved_profile.mode,
        "disable_ocr": resolved_profile.disable_ocr,
        "use_llm": resolved_profile.use_llm,
        "min_recon_score": settings.marker_table_min_recon_score,
        "max_concurrency": 1,
        "processors": ",".join(_official_marker_processor_paths(resolved_profile)),
        # El servicio se instancia desde la ruta de clase pública que espera
        # PdfConverter. Este bloque transporta los Settings ya resueltos sin
        # recurrir a variables de proveedor ni a un fallback implícito.
        "rag_marker_service_settings": {
            "marker_openai_base_url": settings.marker_openai_base_url,
            "marker_openai_model": settings.marker_openai_model,
            "surya_base_url": settings.surya_base_url,
            "local_model_hosts": settings.local_model_hosts,
            "marker_llm_max_requests": settings.marker_llm_max_requests,
            "marker_llm_max_tokens_per_request": settings.marker_llm_max_tokens_per_request,
            "marker_llm_max_generated_tokens_per_document": (
                settings.marker_llm_max_generated_tokens_per_document
            ),
            "marker_llm_request_timeout": settings.marker_llm_request_timeout,
            "marker_llm_job_timeout": settings.marker_llm_job_timeout,
            "marker_llm_max_retries": settings.marker_llm_max_retries,
            "marker_llm_fallback_to_base": settings.marker_llm_fallback_to_base,
            "model_health_connect_timeout": settings.model_health_connect_timeout,
        },
    }

    if settings.marker_strip_existing_ocr:
        config["strip_existing_ocr"] = True

    if settings.marker_page_range.strip():
        config["page_range"] = settings.marker_page_range.strip()

    return config


def _require_marker_2() -> str:
    """Exige Marker para parsear y avisa si no es la versión validada."""
    return require_marker_installed_and_warn_if_unvalidated()


def create_marker_converter(settings: Settings) -> Any:
    """Inicializo Marker 2 una vez y reutilizo sus modelos entre documentos.

    Raises:
        RuntimeError: Si `marker-pdf` no está instalado o no puede inicializarse.
    """
    profile = resolve_marker_profile(settings)

    if profile.use_llm:
        # Se valida antes de crear Marker: ConfigParser nunca puede seleccionar
        # Gemini ni otro proveedor externo como servicio implícito.
        from rag_cliente.marker_llm import validate_marker_local_only

        validate_marker_local_only(settings)

    _require_marker_2()
    os.environ["TORCH_DEVICE"] = profile.torch_device

    if profile.inference_backend is None:
        os.environ.pop("SURYA_INFERENCE_BACKEND", None)
    else:
        os.environ["SURYA_INFERENCE_BACKEND"] = profile.inference_backend

    llama_cpp_binary = settings.llama_cpp_binary.strip()
    if profile.inference_backend == "llamacpp":
        resolved_binary = Path(llama_cpp_binary).expanduser().resolve()
        if not resolved_binary.is_file():
            raise RuntimeError(
                f"LLAMA_CPP_BINARY no existe o no es un archivo: {resolved_binary}"
            )
        os.environ["LLAMA_CPP_BINARY"] = str(resolved_binary)
        os.environ["SURYA_GGUF_LOCAL_MODEL_PATH"] = str(
            Path(settings.surya_gguf_path).expanduser().resolve()
            if settings.surya_gguf_path.strip()
            else (settings.models_path / "surya-ocr-2" / "surya-2.gguf").resolve()
        )
        os.environ["SURYA_GGUF_LOCAL_MMPROJ_PATH"] = str(
            Path(settings.surya_mmproj_path).expanduser().resolve()
            if settings.surya_mmproj_path.strip()
            else (settings.models_path / "surya-ocr-2" / "surya-2-mmproj.gguf").resolve()
        )
        # SURYA_INFERENCE_URL fuerza a Surya a adjuntarse al servidor local ya
        # administrado; así no crea PIDs ni descarga GGUF por su cuenta.
        os.environ["SURYA_INFERENCE_URL"] = settings.surya_base_url
        os.environ["SURYA_INFERENCE_PARALLEL"] = "1"
        os.environ["SURYA_INFERENCE_CTX_SIZE"] = str(settings.model_context_size)
        os.environ["SURYA_INFERENCE_STARTUP_TIMEOUT"] = str(settings.model_start_timeout)
        os.environ["SURYA_INFERENCE_TIMEOUT_SECONDS"] = str(settings.model_request_timeout)

    if profile.torch_device == "cuda":
        try:
            import torch
        except ImportError as exc:
            raise RuntimeError(
                "PyTorch no esta instalado. Ejecuta setup.ps1 para instalar la version CUDA."
            ) from exc

        if not torch.cuda.is_available():
            raise RuntimeError(
                "El perfil gpu-quality requiere CUDA, pero PyTorch no detecta una GPU NVIDIA. "
                "Ejecuta setup.ps1 -Device cuda y comprueba el resultado con 'rag.bat gpu'."
            )

    try:
        from marker.config.parser import ConfigParser
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
    except ImportError as exc:
        raise RuntimeError(
            "Marker 2 no está instalado con todos los proveedores. "
            "Ejecuta setup.ps1 para instalar marker-pdf[full]."
        ) from exc

    try:
        config_parser = ConfigParser(_build_marker_config(settings, profile))
        renderer = (
            "rag_cliente.pdf_loader.MarkerMarkdownRendererWithOcrMetadata"
            if settings.marker_markdown_compatibility
            else "rag_cliente.pdf_loader.MarkerJSONRendererWithOcrMetadata"
        )
        llm_service_path = (
            "rag_cliente.marker_llm.BudgetedMarkerOpenAIService"
            if profile.use_llm
            else None
        )

        converter = PdfConverter(
            config=config_parser.generate_config_dict(),
            artifact_dict=create_model_dict(
                inference_backend=profile.inference_backend
            ),
            processor_list=config_parser.get_processors(),
            renderer=renderer,
            llm_service=llm_service_path,
        )
        if profile.use_llm:
            # PdfConverter publica la instancia que resolvió desde la ruta
            # anterior. Se reutiliza para abrir/cerrar el presupuesto por
            # documento sin tocar internals ni procesadores.
            setattr(converter, "_rag_llm_service", converter.llm_service)
        return converter
    except Exception as exc:
        raise RuntimeError(f"No se pudo inicializar Marker: {exc}") from exc


def _render_with_marker(path: Path, marker_converter: Any) -> tuple[Any, dict[str, Any]]:
    """Ejecuta el converter recibido y conserva su salida oficial completa."""
    llm_service = getattr(marker_converter, "_rag_llm_service", None)
    if llm_service is None:
        rendered = marker_converter(str(path))
    else:
        with llm_service.document_budget(str(path.resolve())):
            rendered = marker_converter(str(path))
    metadata = getattr(rendered, "metadata", {})
    if isinstance(rendered, dict):
        metadata = rendered.get("metadata", metadata)
    if not isinstance(metadata, dict):
        metadata = {}
    return rendered, _metadata_with_marker_capabilities(metadata)


def _marker_markdown_text(rendered: Any) -> str:
    """Obtiene texto del renderer Markdown oficial para el flag de compatibilidad."""
    try:
        from marker.output import text_from_rendered
    except ImportError as exc:
        raise RuntimeError(
            "Marker está instalado parcialmente o cambió su API: no se pudo importar "
            "marker.output.text_from_rendered. Reinstala/actualiza marker-pdf."
        ) from exc

    text, _, _images = text_from_rendered(rendered)
    return text if isinstance(text, str) else ""


def _to_plain_data(value: Any) -> Any:
    """Convierte modelos Pydantic de Marker a tipos JSON sin conocer sus clases."""
    if hasattr(value, "model_dump"):
        try:
            dumped_value = value.model_dump(mode="json")
        except TypeError:
            dumped_value = value.model_dump()
        return _to_plain_data(dumped_value)
    if isinstance(value, dict):
        return {str(key): _to_plain_data(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_to_plain_data(item) for item in value]
    return value


def _block_page_number(block: dict[str, Any], inherited_page: int | None = None) -> int:
    """Obtiene página humana 1-based desde campos o IDs públicos de Marker."""
    for key in ("page_id", "page"):
        if key in block:
            page_number = _as_1_based_page_number(block.get(key))
            if page_number is not None:
                return page_number

    block_id = str(block.get("id", ""))
    match = re.search(r"(?:^|/)page/(\d+)(?:/|$)", block_id, flags=re.IGNORECASE)
    if match:
        return int(match.group(1)) + 1
    return inherited_page or 1


def _text_from_marker_block(block: dict[str, Any]) -> str:
    """Deriva texto indexable sin perder la estructura JSON original."""
    direct_text = block.get("text")
    if isinstance(direct_text, str) and direct_text.strip():
        return _normalize_text(direct_text)

    direct_html = block.get("html")
    if isinstance(direct_html, str) and direct_html.strip():
        without_tags = re.sub(r"<[^>]+>", " ", direct_html)
        normalized_html_text = re.sub(
            r"[ \t]+",
            " ",
            html_lib.unescape(without_tags),
        )
        normalized_html_text = _normalize_text(normalized_html_text)
        if normalized_html_text:
            return normalized_html_text

    children = block.get("children")
    child_texts = [
        _text_from_marker_block(child)
        for child in (children if isinstance(children, list) else [])
        if isinstance(child, dict)
    ]
    return _normalize_text("\n\n".join(text for text in child_texts if text))


def _page_extraction_metadata(metadata: dict[str, Any], page_number: int) -> dict[str, Any]:
    """Conserva metadata general y la estadística correspondiente a la página."""
    page_stats = metadata.get("page_stats")
    selected_page_stats: list[dict[str, Any]] = []
    if isinstance(page_stats, list):
        for page_stat in page_stats:
            if not isinstance(page_stat, dict):
                continue
            if _as_1_based_page_number(page_stat.get("page_id")) == page_number:
                selected_page_stats.append(dict(page_stat))

    extraction_metadata = {
        key: _to_plain_data(value)
        for key, value in metadata.items()
        if key != "page_stats"
    }
    if selected_page_stats:
        extraction_metadata["page_stats"] = selected_page_stats
    return extraction_metadata


def _marker_polygon(block: dict[str, Any]) -> list[list[float]]:
    polygon = _to_plain_data(block.get("polygon") or [])
    if isinstance(polygon, dict):
        polygon = polygon.get("polygon", [])
    if not isinstance(polygon, list):
        return []
    return [list(point) for point in polygon if isinstance(point, (list, tuple))]


def _marker_confidence(block: dict[str, Any], kind: str) -> float | None:
    """Conserva confianza explícita sin inventar umbrales ni decisiones."""
    confidence = block.get("confidence")
    if isinstance(confidence, (int, float)) and not isinstance(confidence, bool):
        return float(confidence)

    top_k = block.get("top_k")
    if isinstance(top_k, dict):
        for label, value in top_k.items():
            if str(label).lower() != kind.lower():
                continue
            if isinstance(value, (int, float)) and not isinstance(value, bool):
                return float(value)
    return None


def _deduplicate(values: list[Any]) -> list[Any]:
    result: list[Any] = []
    seen: set[Any] = set()
    for value in values:
        key = value
        if isinstance(value, dict):
            key = (value.get("page"), value.get("block_id"))
        if key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def _normalize_marker_element(
    block: dict[str, Any],
    *,
    inherited_page: int | None = None,
    parent_id: str | None = None,
    inherited_section_path: list[str] | None = None,
    inherited_table_id: str | None = None,
    reading_order: list[int] | None = None,
) -> DocumentElement:
    """Propaga solo procedencia observable en el JSON oficial de Marker."""
    counter = reading_order if reading_order is not None else [0]
    kind = str(block.get("block_type") or block.get("kind") or "Unknown")
    block_id = str(block.get("id") or "")
    page_number = _block_page_number(block, inherited_page)
    polygon = _marker_polygon(block)
    section_path = _section_path_from_block(block, inherited_section_path)
    normalized_kind = kind.strip().lower()
    table_id = (
        str(block.get("table_id") or block_id or "").strip() or None
        if normalized_kind in {"table", "tableofcontents"}
        else inherited_table_id
    )

    own_span: list[dict[str, Any]] = []
    if block_id and kind.lower() not in {"document", "page"}:
        own_span.append(
            {
                "page": page_number,
                "block_id": block_id,
                "polygon": polygon,
                "reading_order": counter[0],
            }
        )
        counter[0] += 1

    raw_children = block.get("children")
    children: list[DocumentElement] = []
    if isinstance(raw_children, list):
        for child in raw_children:
            if isinstance(child, dict):
                children.append(
                    _normalize_marker_element(
                        child,
                        inherited_page=page_number,
                        parent_id=block_id or parent_id,
                        inherited_section_path=section_path,
                        inherited_table_id=table_id,
                        reading_order=counter,
                    )
                )

    child_spans = [span for child in children for span in child.source_spans]
    source_spans = _deduplicate([*own_span, *child_spans])
    source_pages = sorted(
        {
            int(span["page"])
            for span in source_spans
            if isinstance(span.get("page"), int)
        }
    )
    if not source_pages:
        source_pages = [page_number]
    source_block_ids = _deduplicate(
        [
            str(span["block_id"])
            for span in source_spans
            if str(span.get("block_id") or "")
        ]
    )

    section_hierarchy = _to_plain_data(block.get("section_hierarchy") or {})
    if not isinstance(section_hierarchy, dict):
        section_hierarchy = {}

    return DocumentElement(
        id=block_id,
        kind=kind,
        html=str(block.get("html") or ""),
        text=_text_from_marker_block(block),
        page_start=min(source_pages),
        page_end=max(source_pages),
        source_pages=source_pages,
        source_block_ids=source_block_ids,
        source_spans=source_spans,
        section_hierarchy=section_hierarchy,
        polygon=polygon,
        confidence=_marker_confidence(block, kind),
        children=children,
        parent_id=parent_id,
        table_id=table_id,
        section_path=section_path,
        provenance="marker",
        metadata=(
            dict(_to_plain_data(block.get("metadata")))
            if isinstance(_to_plain_data(block.get("metadata")), dict)
            else {}
        ),
    )


def _extract_marker_structured_chunks(
    rendered: Any,
    metadata: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    """Normaliza la salida JSON oficial al contrato estructurado del RAG."""
    payload = _to_plain_data(rendered)
    if not isinstance(payload, dict):
        return []

    effective_metadata = _metadata_with_marker_capabilities(metadata or {})
    payload_metadata = payload.get("metadata")
    if isinstance(payload_metadata, dict):
        effective_metadata.update(payload_metadata)
    effective_metadata = _metadata_with_marker_capabilities(effective_metadata)

    root_children = payload.get("children")
    candidates = (
        [child for child in root_children if isinstance(child, dict)]
        if isinstance(root_children, list)
        else [payload]
    )

    page_candidates = [
        block
        for block in candidates
        if str(block.get("block_type", "")).strip().lower() == "page"
    ]
    blocks = page_candidates or candidates
    structured_chunks: list[dict[str, Any]] = []

    reading_order = [0]
    for block in blocks:
        element = _normalize_marker_element(block, reading_order=reading_order)
        element_data = element.as_dict()
        structured_chunks.append(
            {
                **element_data,
                # Alias temporales consumidos por PageDocument y código previo.
                "block_type": element.kind,
                "page": element.page_start,
                "extraction_metadata": _page_extraction_metadata(
                    effective_metadata,
                    element.page_start,
                ),
            }
        )

    return structured_chunks


def _extract_marker_ocr_usage_by_page(metadata: dict[str, Any]) -> dict[int, bool]:
    """Extrae uso de OCR por página a partir de metadatos de Marker.

    Prioriza el campo enriquecido `ocr_used`. Si no existe, cae a
    `text_extraction_methods` o `text_extraction_method`. `pdftext` significa
    texto nativo del PDF; `surya`/`gemini`/otros valores implican OCR o visión.
    """
    page_stats = metadata.get("page_stats")
    if not isinstance(page_stats, list):
        return {}

    ocr_usage_by_page: dict[int, bool] = {}
    for page_stat in page_stats:
        if not isinstance(page_stat, dict):
            continue

        page_number = _as_1_based_page_number(page_stat.get("page_id"))
        if page_number is None:
            continue

        if isinstance(page_stat.get("ocr_used"), bool):
            ocr_usage_by_page[page_number] = page_stat["ocr_used"]
            continue

        methods: list[str] = []
        raw_methods = page_stat.get("text_extraction_methods")
        if isinstance(raw_methods, list):
            methods.extend(str(method).strip().lower() for method in raw_methods if str(method).strip())

        method = str(page_stat.get("text_extraction_method", "")).strip().lower()
        if method:
            methods.append(method)

        if methods:
            ocr_usage_by_page[page_number] = any(method != "pdftext" for method in methods)

    return ocr_usage_by_page


def _split_marker_markdown_by_page(markdown: str) -> list[tuple[int, str]]:
    """Divido el Markdown de Marker 2 en pares con página humana 1-based.

    Marker 2 escribe siempre el `page_id` interno 0-based en el separador. Sumo
    uno incluso cuando yo proceso un rango que empieza después de la página 0;
    así evito que las citas queden desplazadas al usar `MARKER_PAGE_RANGE`.
    """
    normalized_markdown = markdown.strip()
    if not normalized_markdown:
        return []

    matches = list(_MARKER_PAGE_SEPARATOR_RE.finditer(normalized_markdown))
    if not matches:
        return [(1, _normalize_text(normalized_markdown))]

    raw_page_numbers = [int(match.group(1)) for match in matches]
    pages: list[tuple[int, str]] = []

    preamble = normalized_markdown[: matches[0].start()].strip()
    if preamble:
        first_page = raw_page_numbers[0] + 1
        pages.append((max(first_page, 1), _normalize_text(preamble)))

    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(normalized_markdown)
        text = _normalize_text(normalized_markdown[start:end])
        if not text:
            continue
        page_number = int(match.group(1)) + 1
        pages.append((max(page_number, 1), text))

    return pages


def _marker_document_pages(
    path: Path,
    source_type: str,
    marker_converter: Any,
    settings: Settings,
    default_ocr_used: bool = False,
) -> list[PageDocument]:
    """Adapta JSON estructurado o el Markdown temporal a `PageDocument`."""
    try:
        rendered, metadata = _render_with_marker(path, marker_converter)
    except Exception as exc:
        from rag_cliente.marker_llm import MarkerLLMError

        if isinstance(exc, MarkerLLMError):
            raise
        raise RuntimeError(f"Marker falló procesando '{path.name}': {exc}") from exc

    marker_ocr_usage = _extract_marker_ocr_usage_by_page(metadata)
    document_ocr_used = any(marker_ocr_usage.values()) if marker_ocr_usage else default_ocr_used

    if settings.marker_markdown_compatibility:
        markdown = _marker_markdown_text(rendered)
        structured_chunks = [
            {
                "block_type": "Page",
                "id": f"/page/{page_number - 1}",
                "html": "",
                "text": text,
                "page": page_number,
                "page_start": page_number,
                "page_end": page_number,
                "source_pages": [page_number],
                "source_block_ids": [],
                "source_spans": [],
                "polygon": [],
                "children": [],
                "section_hierarchy": {},
                "confidence": None,
                "extraction_metadata": _page_extraction_metadata(metadata, page_number),
            }
            for page_number, text in _split_marker_markdown_by_page(markdown)
        ]
    else:
        structured_chunks = _extract_marker_structured_chunks(rendered, metadata)

    pages: list[PageDocument] = []
    parser_profile = resolve_marker_profile(settings).name
    for chunk in structured_chunks:
        page_number = int(chunk["page"])
        text = str(chunk["text"])
        if not text:
            continue
        ocr_used = marker_ocr_usage.get(
            page_number,
            document_ocr_used if len(structured_chunks) == 1 else default_ocr_used,
        )
        pages.append(
            PageDocument(
                document_id=path.stem,
                source=path.name,
                source_path=str(path.resolve()),
                source_type=source_type,
                page_number=page_number,
                text=text,
                ocr_used=ocr_used,
                block_type=str(chunk["block_type"]),
                id=str(chunk["id"]),
                html=str(chunk["html"]),
                page=page_number,
                polygon=chunk["polygon"],
                children=chunk["children"],
                section_hierarchy=chunk["section_hierarchy"],
                page_start=int(chunk.get("page_start", page_number)),
                page_end=int(chunk.get("page_end", page_number)),
                source_pages=[int(page) for page in chunk.get("source_pages", [page_number])],
                source_block_ids=[
                    str(block_id) for block_id in chunk.get("source_block_ids", [])
                ],
                source_spans=[
                    dict(span)
                    for span in chunk.get("source_spans", [])
                    if isinstance(span, dict)
                ],
                confidence=chunk.get("confidence"),
                extraction_metadata=chunk["extraction_metadata"],
                provenance="marker",
                parser_profile=parser_profile,
            )
        )
    return pages


def _document_element_from_mapping(
    value: dict[str, Any],
    *,
    document_id: str,
    parent_id: str | None = None,
    inherited_table_id: str | None = None,
    inherited_section_path: list[str] | None = None,
) -> DocumentElement:
    kind = str(value.get("kind") or value.get("block_type") or "Unknown")
    element_id = str(value.get("id") or "")
    normalized_kind = kind.strip().lower()
    table_id = (
        str(value.get("table_id") or element_id or "").strip() or None
        if normalized_kind in {"table", "tableofcontents"}
        else inherited_table_id
    )
    section_path = [
        str(item)
        for item in value.get("section_path", inherited_section_path or [])
        if str(item).strip()
    ]
    children = [
        _document_element_from_mapping(
            child,
            document_id=document_id,
            parent_id=element_id or parent_id,
            inherited_table_id=table_id,
            inherited_section_path=section_path,
        )
        for child in value.get("children", [])
        if isinstance(child, dict)
    ]
    page_start = int(value.get("page_start", value.get("page", 1)))
    page_end = int(value.get("page_end", page_start))
    polygon = value.get("polygon", [])
    return DocumentElement(
        id=element_id,
        kind=kind,
        html=str(value.get("html") or ""),
        text=str(value.get("text") or ""),
        page_start=page_start,
        page_end=page_end,
        source_pages=[int(page) for page in value.get("source_pages", [page_start])],
        source_block_ids=[str(item) for item in value.get("source_block_ids", [])],
        source_spans=[dict(span) for span in value.get("source_spans", [])],
        section_hierarchy=dict(value.get("section_hierarchy") or {}),
        polygon=[list(point) for point in polygon if isinstance(point, (list, tuple))],
        confidence=value.get("confidence"),
        children=children,
        document_id=document_id,
        parent_id=(
            str(value.get("parent_id")).strip()
            if value.get("parent_id") is not None
            else parent_id
        ),
        table_id=table_id,
        section_path=section_path,
        provenance=str(value.get("provenance") or "marker"),
        metadata=dict(value.get("metadata") or {}),
    )


def parsed_document_from_pages(path: Path, pages: list[PageDocument]) -> ParsedDocument:
    """Agrupa PageDocument sin alterar ni inferir la estructura de Marker."""
    elements = [
        DocumentElement(
            id=page.id,
            kind=page.block_type or "Page",
            html=page.html,
            text=page.text,
            page_start=int(page.page_start or page.page_number),
            page_end=int(page.page_end or page.page_number),
            source_pages=list(page.source_pages),
            source_block_ids=list(page.source_block_ids),
            source_spans=[dict(span) for span in page.source_spans],
            section_hierarchy=dict(page.section_hierarchy),
            polygon=[list(point) for point in page.polygon],
            confidence=page.confidence,
            children=[
                _document_element_from_mapping(
                    child,
                    document_id=path.stem,
                    parent_id=page.id or None,
                    inherited_section_path=[
                        str(value)
                        for _, value in sorted(page.section_hierarchy.items())
                        if str(value).strip()
                    ],
                )
                for child in page.children
                if isinstance(child, dict)
            ],
            document_id=path.stem,
            table_id=(
                page.id or None
                if (page.block_type or "").strip().lower() in {"table", "tableofcontents"}
                else None
            ),
            section_path=[
                str(value)
                for _, value in sorted(page.section_hierarchy.items())
                if str(value).strip()
            ],
            provenance=page.provenance,
        )
        for page in pages
    ]

    metadata: dict[str, Any] = {
        "capabilities": marker_capabilities(),
        "tag": next((page.tag for page in pages if page.tag), ""),
        "parser_profile": next(
            (page.parser_profile for page in pages if page.parser_profile),
            "unknown",
        ),
        "ocr_used_by_page": {
            str(page.page_number): bool(page.ocr_used)
            for page in pages
        },
    }
    page_stats: list[dict[str, Any]] = []
    seen_page_ids: set[Any] = set()
    for page in pages:
        for key, value in page.extraction_metadata.items():
            if key != "page_stats":
                metadata.setdefault(key, _to_plain_data(value))
        for page_stat in page.extraction_metadata.get("page_stats", []) or []:
            if not isinstance(page_stat, dict):
                continue
            page_id = page_stat.get("page_id")
            if page_id in seen_page_ids:
                continue
            seen_page_ids.add(page_id)
            page_stats.append(dict(page_stat))
    if page_stats:
        metadata["page_stats"] = page_stats

    return ParsedDocument(
        id=path.stem,
        source=path.name,
        source_path=str(path.resolve()),
        source_type=path.suffix.lower().lstrip("."),
        elements=elements,
        metadata=metadata,
    )


def parsed_documents_from_pages(pages: list[PageDocument]) -> list[ParsedDocument]:
    """Agrupa unidades por documento y expone la estructura normalizada completa."""
    grouped: dict[str, list[PageDocument]] = {}
    order: list[str] = []
    for page in pages:
        key = page.source_path or f"{page.document_id}:{page.source}"
        if key not in grouped:
            grouped[key] = []
            order.append(key)
        grouped[key].append(page)

    documents: list[ParsedDocument] = []
    for key in order:
        document_pages = grouped[key]
        source_path = Path(document_pages[0].source_path or document_pages[0].source)
        documents.append(parsed_document_from_pages(source_path, document_pages))
    return documents


def load_marker_document_pages(
    document_path: Path,
    settings: Settings,
    marker_converter: Any | None = None,
) -> list[PageDocument]:
    """Proceso con Marker 2 cualquier formato admitido por su paquete full.

    Dejo que el modo de Marker decida la estrategia para cada página y bloque.
    Solo marco OCR por defecto en imágenes puras o cuando yo lo fuerzo de forma
    explícita; para el resto uso los metadatos reales del renderer enriquecido.
    """
    converter = marker_converter or create_marker_converter(settings)
    suffix = document_path.suffix.lower()
    profile = resolve_marker_profile(settings)
    return _marker_document_pages(
        document_path,
        source_type=suffix.lstrip("."),
        marker_converter=converter,
        settings=settings,
        default_ocr_used=suffix in IMAGE_SUFFIXES and not profile.disable_ocr,
    )


def load_pdf_pages(
    pdf_path: Path,
    settings: Settings | None = None,
    ocr_pipeline: Any | None = None,
    marker_converter: Any | None = None,
) -> list[PageDocument]:
    """Cargo un PDF con Marker 2 o con mi fallback digital si lo desactivo.

    `ocr_pipeline` se mantiene por compatibilidad con llamadas antiguas del
    backend; si se recibe, se trata como un converter de Marker ya inicializado.
    """
    if settings is None or not settings.marker_enabled:
        return _load_native_pdf_pages(pdf_path)

    return load_marker_document_pages(
        pdf_path,
        settings=settings,
        marker_converter=marker_converter or ocr_pipeline,
    )


def load_docx_pages(docx_path: Path) -> list[PageDocument]:
    """Carga y extrae texto de un documento Word (.docx)."""
    document = Document(docx_path)
    text = _extract_docx_text(document)
    if not text:
        return []

    return [
        PageDocument(
            document_id=docx_path.stem,
            source=docx_path.name,
            source_path=str(docx_path.resolve()),
            source_type="docx",
            page_number=1,
            text=text,
            ocr_used=False,
            provenance="python-docx",
            parser_profile="native-docx",
        )
    ]


def load_txt_pages(txt_path: Path) -> list[PageDocument]:
    """Carga y extrae texto de un archivo de texto plano UTF-8."""
    text = txt_path.read_text(encoding="utf-8").strip()
    if not text:
        return []

    return [
        PageDocument(
            document_id=txt_path.stem,
            source=txt_path.name,
            source_path=str(txt_path.resolve()),
            source_type="txt",
            page_number=1,
            text=_normalize_text(text),
            ocr_used=False,
            provenance="text",
            parser_profile="native-text",
        )
    ]


def load_image_pages(
    image_path: Path,
    settings: Settings | None = None,
    ocr_pipeline: Any | None = None,
    marker_converter: Any | None = None,
) -> list[PageDocument]:
    """Cargo una imagen mediante el OCR adaptativo de Marker 2.

    Mantiene compatibilidad funcional con el soporte previo de imágenes sin
    depender de PaddleOCR/PaddleX.
    """
    if settings is None or not settings.marker_enabled:
        return []

    return load_marker_document_pages(
        image_path,
        settings=settings,
        marker_converter=marker_converter or ocr_pipeline,
    )


def load_documents_from_directory(
    doc_dir: Path,
    settings: Settings | None = None,
    progress_callback: ProgressCallback | None = None,
) -> list[PageDocument]:
    """Cargo recursivamente todos los formatos soportados por Marker 2 full.

    Si un archivo falla, continúo con los demás y dejo el motivo visible en el
    progreso para que yo pueda diagnosticarlo sin perder toda la indexación.
    """
    if not doc_dir.exists():
        raise FileNotFoundError(f"Document directory does not exist: {doc_dir}")

    base_dir = doc_dir.resolve()
    paths = [path for path in sorted(doc_dir.rglob("*")) if path.is_file()]
    supported_paths = [
        path
        for path in paths
        if path.suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES
    ]

    marker_converter: Any | None = None
    if settings is not None and settings.marker_enabled:
        needs_marker = any(path.suffix.lower() in MARKER_DOCUMENT_SUFFIXES for path in supported_paths)
        if needs_marker:
            _emit(progress_callback, "Inicializando Marker...")
            marker_converter = create_marker_converter(settings)

    all_pages: list[PageDocument] = []
    total_files = len(supported_paths)

    for file_index, path in enumerate(supported_paths, start=1):
        suffix = path.suffix.lower()
        try:
            relative_path = path.resolve().relative_to(base_dir)
        except ValueError:
            relative_path = Path(path.name)
        tag = relative_path.parts[0] if len(relative_path.parts) > 1 else ""
        display_path = str(relative_path)
        _emit(progress_callback, f"Procesando archivo {file_index}/{total_files}: {display_path}")

        try:
            if suffix == ".txt":
                pages = load_txt_pages(path)
            elif settings is not None and settings.marker_enabled and suffix in MARKER_DOCUMENT_SUFFIXES:
                _emit(
                    progress_callback,
                    f"Parseando {suffix.lstrip('.').upper()} con Marker 2 "
                    f"({resolve_marker_profile(settings).name}): {display_path}",
                )
                pages = load_marker_document_pages(
                    path,
                    settings=settings,
                    marker_converter=marker_converter,
                )
            elif suffix == ".pdf":
                pages = _load_native_pdf_pages(path)
            elif suffix == ".docx":
                pages = load_docx_pages(path)
            else:
                continue
        except Exception as exc:
            from rag_cliente.marker_llm import MarkerLLMError

            if isinstance(exc, MarkerLLMError):
                raise
            _emit(progress_callback, f"AVISO: no se pudo procesar {display_path}: {exc}")
            continue

        if pages:
            for page in pages:
                page.tag = tag
            all_pages.extend(pages)
            ocr_pages = sum(1 for page in pages if page.ocr_used)
            tag_label = f", tag: {tag}" if tag else ""
            _emit(
                progress_callback,
                f"Extraídas {len(pages)} páginas/bloques de {display_path} "
                f"(OCR usado en {ocr_pages}{tag_label})",
            )
        else:
            _emit(progress_callback, f"AVISO: {display_path} no produjo texto indexable")

    return all_pages
